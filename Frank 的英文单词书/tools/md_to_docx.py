#!/usr/bin/env python3
"""Convert Frank vocabulary markdown to print-ready DOCX with pronunciation links."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor

from pronunciation import fetch_pronunciation
from vocab_md import PHONETIC_RE, lookup_word, parse_markdown


def parse_bold_runs(text: str, paragraph, *, base_bold: bool = False, base_italic: bool = False):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.italic = base_italic
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic


def set_run_font(run, *, western: str, east_asia: str, size_pt: float, bold=False, italic=False, color=None):
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_hyperlink(paragraph, url: str, text: str, *, size_pt: float = 10, color_hex: str = "1A73E8"):
    if not url:
        return
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    r_pr.append(sz)
    new_run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run("1")
    run5 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_end)


def clear_paragraph(paragraph):
    paragraph.text = ""


def configure_headers_footers(doc: Document, chapter_title: str | None):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    clear_paragraph(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if chapter_title:
        run = hp.add_run(f"Frank 的英文单词书 · {chapter_title}")
        set_run_font(run, western="Times New Roman", east_asia="Songti SC", size_pt=9, color=RGBColor(0x66, 0x66, 0x66))

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)
    for run in fp.runs:
        set_run_font(run, western="Times New Roman", east_asia="Songti SC", size_pt=9, color=RGBColor(0x66, 0x66, 0x66))


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(17.6)
    section.page_height = Cm(25.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.0)
    section.gutter = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.25
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def add_chapter_title(doc: Document, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(title)
    set_run_font(run, western="Times New Roman", east_asia="Songti SC", size_pt=18, bold=True)


def add_headword(doc: Document, word: str, *, first: bool):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10 if first else 14)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    run = p.add_run(word)
    set_run_font(run, western="Georgia", east_asia="Georgia", size_pt=13.5, bold=True, color=RGBColor(0x1A, 0x1A, 0x1A))


def _speaker_phonetic_index(meta_lines: list[str]) -> int | None:
    indices = [i for i, line in enumerate(meta_lines) if PHONETIC_RE.match(line)]
    if not indices:
        return None
    for i in indices:
        if meta_lines[i].startswith("美"):
            return i
    return indices[0]


def add_meta_line(
    doc: Document,
    text: str,
    *,
    italic: bool = False,
    indent: bool = False,
    pronunciation=None,
    show_speaker: bool = False,
):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15
    if indent:
        pf.left_indent = Cm(0.35)

    parse_bold_runs(text, p, base_italic=italic)
    for run in p.runs:
        size = 10 if italic else 10.5
        set_run_font(
            run,
            western="Times New Roman",
            east_asia="Songti SC",
            size_pt=size,
            bold=run.bold,
            italic=run.italic or italic,
            color=RGBColor(0x33, 0x33, 0x33) if italic else None,
        )

    if show_speaker and pronunciation and pronunciation.audio:
        spacer = p.add_run("  ")
        set_run_font(spacer, western="Times New Roman", east_asia="Songti SC", size_pt=10)
        add_hyperlink(p, pronunciation.audio, "🔊", size_pt=10, color_hex="1A73E8")


def add_field_line(doc: Document, label: str, content: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.2
    pf.left_indent = Cm(0.15)
    pf.first_line_indent = Cm(-0.15)

    label_run = p.add_run(f"[{label}] ")
    set_run_font(
        label_run,
        western="Times New Roman",
        east_asia="Heiti SC",
        size_pt=10,
        bold=True,
        color=RGBColor(0x2E, 0x4A, 0x6E),
    )
    parse_bold_runs(content, p)
    for run in p.runs[1:]:
        set_run_font(
            run,
            western="Times New Roman",
            east_asia="Songti SC",
            size_pt=10.5,
            bold=run.bold,
            italic=run.italic,
        )


def build_docx(chapter_title: str | None, entries: list[dict], *, with_pronunciation: bool = True) -> Document:
    doc = Document()
    configure_document(doc)
    configure_headers_footers(doc, chapter_title)

    if chapter_title:
        add_chapter_title(doc, chapter_title)

    for idx, entry in enumerate(entries):
        add_headword(doc, entry["word"], first=idx == 0)

        pronunciation = None
        if with_pronunciation:
            pronunciation = fetch_pronunciation(lookup_word(entry["word"]))

        speaker_idx = _speaker_phonetic_index(entry["meta"])
        for i, meta in enumerate(entry["meta"]):
            is_phonetic = bool(PHONETIC_RE.match(meta))
            add_meta_line(
                doc,
                meta,
                italic=is_phonetic,
                show_speaker=i == speaker_idx,
                pronunciation=pronunciation,
            )

        for label, content in entry["fields"]:
            add_field_line(doc, label, content)

    return doc


def convert(md_path: Path, out_path: Path | None = None, *, with_pronunciation: bool = True) -> Path:
    text = md_path.read_text(encoding="utf-8")
    chapter_title, entries = parse_markdown(text)
    if not entries:
        raise SystemExit(f"No entries found in {md_path}")

    doc = build_docx(chapter_title, entries, with_pronunciation=with_pronunciation)
    if out_path is None:
        out_path = md_path.with_suffix(".docx")
    doc.save(out_path)
    return out_path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Convert vocabulary markdown to print-ready DOCX")
    parser.add_argument("input", type=Path, help="Input .md file")
    parser.add_argument("-o", "--output", type=Path, help="Output .docx path")
    parser.add_argument("--no-pronunciation", action="store_true", help="Skip pronunciation link lookup")
    args = parser.parse_args(argv)

    if not args.input.is_file():
        print(f"Error: file not found: {args.input}", file=sys.stderr)
        return 1

    out = convert(args.input, args.output, with_pronunciation=not args.no_pronunciation)
    print(f"Wrote {out} ({out.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
